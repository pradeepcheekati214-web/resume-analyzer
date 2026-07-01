"""
Resume text extraction from PDF and DOCX files.
"""
import io
import logging
import re
from dataclasses import dataclass, field
from typing import Optional

logger = logging.getLogger(__name__)


@dataclass
class ParsedResume:
    raw_text: str = ""
    word_count: int = 0
    page_count: int = 0
    file_type: str = ""
    contact_info: dict = field(default_factory=dict)


def parse_resume(file_bytes: bytes, filename: str) -> ParsedResume:
    """
    Extract text and metadata from a resume file.
    Supports PDF and DOCX. Falls back to plain-text if parsing fails.
    """
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else "pdf"
    if ext == "pdf":
        return _parse_pdf(file_bytes)
    elif ext in ("docx", "doc"):
        return _parse_docx(file_bytes)
    else:
        return _parse_plaintext(file_bytes, ext)


# ---------------------------------------------------------------------------
# PDF parsing
# ---------------------------------------------------------------------------
def _parse_pdf(file_bytes: bytes) -> ParsedResume:
    raw_text = ""
    page_count = 1

    # Try pdfplumber first
    try:
        import pdfplumber
        text_pages = []
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            page_count = len(pdf.pages)
            for page in pdf.pages:
                page_text = page.extract_text() or ""
                text_pages.append(page_text)
        raw_text = "\n".join(text_pages)
    except Exception as e1:
        logger.warning("pdfplumber failed: %s — trying PyPDF2", e1)
        # Try PyPDF2 fallback
        try:
            from PyPDF2 import PdfReader
            reader = PdfReader(io.BytesIO(file_bytes))
            pages = [p.extract_text() or "" for p in reader.pages]
            raw_text = "\n".join(pages)
            page_count = len(pages)
        except Exception as e2:
            logger.warning("PyPDF2 failed: %s — trying plain text fallback", e2)
            # Last resort: treat bytes as plain text
            raw_text = _decode_bytes(file_bytes)

    raw_text = _clean_text(raw_text)
    return ParsedResume(
        raw_text=raw_text,
        word_count=len(raw_text.split()),
        page_count=page_count,
        file_type="pdf",
        contact_info=extract_contact_info(raw_text),
    )


# ---------------------------------------------------------------------------
# DOCX parsing
# ---------------------------------------------------------------------------
def _parse_docx(file_bytes: bytes) -> ParsedResume:
    try:
        from docx import Document
        doc = Document(io.BytesIO(file_bytes))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text.strip())
        raw_text = _clean_text("\n".join(paragraphs))
    except Exception as e:
        logger.warning("DOCX parsing failed: %s — trying plain text", e)
        raw_text = _clean_text(_decode_bytes(file_bytes))

    return ParsedResume(
        raw_text=raw_text,
        word_count=len(raw_text.split()),
        page_count=1,
        file_type="docx",
        contact_info=extract_contact_info(raw_text),
    )


# ---------------------------------------------------------------------------
# Plain text fallback
# ---------------------------------------------------------------------------
def _parse_plaintext(file_bytes: bytes, ext: str = "txt") -> ParsedResume:
    raw_text = _clean_text(_decode_bytes(file_bytes))
    return ParsedResume(
        raw_text=raw_text,
        word_count=len(raw_text.split()),
        page_count=1,
        file_type=ext,
        contact_info=extract_contact_info(raw_text),
    )


def _decode_bytes(file_bytes: bytes) -> str:
    """Try UTF-8 then latin-1 to decode bytes to string."""
    for enc in ("utf-8", "latin-1", "cp1252"):
        try:
            return file_bytes.decode(enc)
        except Exception:
            continue
    return file_bytes.decode("utf-8", errors="replace")


# ---------------------------------------------------------------------------
# Contact information extraction
# ---------------------------------------------------------------------------
def extract_contact_info(text: str) -> dict:
    info = {}

    email_match = re.search(r"[\w.+-]+@[\w-]+\.[a-zA-Z]{2,}", text)
    if email_match:
        info["email"] = email_match.group(0)

    phone_match = re.search(r"(\+?[\d][\d\s\-().]{7,}\d)", text)
    if phone_match:
        candidate = phone_match.group(0).strip()
        if len(re.sub(r"\D", "", candidate)) >= 7:
            info["phone"] = candidate

    linkedin = re.search(r"linkedin\.com/in/[\w-]+", text, re.IGNORECASE)
    if linkedin:
        info["linkedin"] = linkedin.group(0)

    github = re.search(r"github\.com/[\w-]+", text, re.IGNORECASE)
    if github:
        info["github"] = github.group(0)

    lines = [l.strip() for l in text.split("\n") if l.strip()]
    for line in lines[:5]:
        words = line.split()
        if 1 < len(words) <= 5 and all(w.replace("-", "").isalpha() for w in words):
            info["name"] = line
            break

    return info


def _clean_text(text: str) -> str:
    text = re.sub(r"[^\x20-\x7E\n]", " ", text)
    text = re.sub(r"[ \t]{2,}", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()
